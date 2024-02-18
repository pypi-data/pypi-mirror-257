from . import Markup
from . import Fonts
from . import SectionMetadata
from . import Metadata
from . import Section
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import traceback
from typing import List, Optional, cast
from . import Tokenizer


class RegularSection(Section):
    def __init__(
        self,
        lines: List[str],
        markup: Markup,
        fonts: Fonts,
        document: Document,
        metadata: Metadata,
    ) -> None:
        super().__init__(lines, markup, fonts, document, metadata)
        # self._quote: Optional[List[str]] = None

    # ============= PUBLIC STUFF HERE

    def compile(self) -> bool:
        """returns True if the lines were all added to the document"""
        reqs = [self._markup, self._fonts, self._document, self._metadata]
        if None in reqs:
            raise Exception(f"Requirements not setup: {reqs}")
        line_number = 0
        for line in self._lines:
            try:
                self._append_line(self._lines, line, line_number)
                line_number = line_number + 1
                # self._lines_count = self._lines_count + 1
            except Exception as e:
                print(f"Error in section at {line_number}: {line}: {e}")
                return False
        self.metadata.FILE_LINE_COUNT = line_number
        return True

    @property
    def doc(self) -> Document:
        return self._document

    @property
    def metadata(self) -> SectionMetadata:
        return cast(SectionMetadata, self._metadata)

    # ============= INTERNAL STUFF STARTS HERE

    def _append_quote(self) -> None:
        """adds the quote into the document"""
        if self._quote is None:  # type: ignore [has-type]
            raise Exception("No quote array found at _append_quote")
        i = len(self._quote)
        for aline in self._quote:
            if aline is None:  # make mypy happy
                continue
            p = self.doc.add_paragraph()
            paragraph_format = p.paragraph_format
            run = self._add_run(p, f"   {aline[1:]}")
            run.font.name = self._fonts.QUOTE
            run.italic = True
            i = i - 1
            if i > 0:
                paragraph_format.space_after = Pt(1)
        self._quote = None

    def _append_block(self) -> None:
        """adds the block text to the document"""
        if not self._block:
            raise Exception("No block array found at _append_block")
        i = len(self._block)
        for aline in self._block:
            if aline is None:  # make mypy happy
                continue
            p = self.doc.add_paragraph()
            paragraph_format = p.paragraph_format
            thisline = aline[1:]
            run = self._add_run(p, f"{thisline}")
            run.italic = True
            run.font.name = self._fonts.BLOCK
            run.font.size = Pt(10)
            i = i - 1
            if i > 0:
                paragraph_format.space_after = Pt(1)
        self._block = None

    def _append_output(self, lines: List[str], line: str, line_number: int) -> None:
        """writes blocks, quotes, and page breaks to the document"""
        if self._block is not None:
            # write the block
            self._append_block()
        if self._quote is not None:
            # write the quote
            self._append_quote()
        if self._part_break and self._last_line(lines, line_number):
            p = self.doc.add_paragraph()
            run = self._add_run(p, "")
            run.font.name = self._fonts.BODY
            run.add_break(WD_BREAK.PAGE)
            self._part_break = False
            self._book_break = False
            self._last_was_break = True

    def _last_line(self, lines: List[str], line_number: int) -> bool:
        """returns True if line_number indicates the last non-blank line"""
        n = len(lines)
        if line_number == n:
            return True
        for r in range(line_number + 1, n):
            if lines[r].strip() != "":
                return False
        return True

    def _append_title(self, line: str) -> None:
        """adds titles to the document"""
        self._part_break = len(line) >= 1 and line[0:1] == self._markup.CHAPTER_TITLE
        self._book_break = len(line) >= 2 and line[0:2] == self._markup.BOOK_TITLE
        simple_separator = len(line) >= 3 and line[0:3] == self._markup.JUMP
        close_part = len(line) >= 1 and line[0:1] == self._markup.NEW_SECTION
        #
        # if we're a _part_break or book_break we need a page break
        #   except if the last file was also a break. in that case
        #   we don't need another blank page
        #
        if (self._part_break and not self._last_was_break) or close_part:
            line = line[2 if self._book_break else 1 :]
            p = self.doc.add_paragraph()
            run = self._add_run(p, "")
            run.font.name = self._fonts.TITLE
            run.add_break(WD_BREAK.PAGE)
        elif self._part_break:
            line = line[2 if self._book_break else 1 :]
        p = None
        if simple_separator:
            #
            # if we're not a _part_break we might be a simple jump separator
            #
            p = self.doc.add_paragraph()
            paragraph_format = p.paragraph_format
            paragraph_format.space_before = Pt(24)
            paragraph_format.space_after = Pt(25)
            run = self._add_jump(p)
        else:
            #
            # we're not a simple separator, therefore we must be a heading of some kind
            #
            if self._part_break:
                p = self.doc.add_heading("", 1 if self._book_break else 2)
            else:
                p = self.doc.add_heading("", 3)
                paragraph_format = p.paragraph_format
                paragraph_format.space_before = Pt(30)
                paragraph_format.space_after = Pt(10)
        if self._part_break or simple_separator:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not simple_separator:
            run = self._add_run(p, line)
            run.font.name = self._fonts.BODY
            run.font.color.rgb = RGBColor.from_string("000000")

    def _handle_block(self, line, line_number, lines: int) -> Optional[bool]:
        block = self._markup._is_block(line, line_number, lines)
        if block:
            if self._block is None:
                self._block = []
            self._block.append(line)
            # pack the line into self.block
            return True
        elif block is None:
            return None
        elif self._block is not None:
            # write the block
            self._append_block()
            return False
        else:
            return False

    def _handle_quote(self, line: str, line_number: int, lines: int) -> Optional[bool]:
        quote = self._markup._is_quote(line, line_number, lines)
        if quote:
            if self._quote is None:
                self._quote = []
            self._quote.append(line)
            # we pack the line into self.quote
            return True
        elif self._quote is not None:
            # write the quote
            self._append_quote()
            return False
        else:
            return False

    def _handle_highlights(self, line):
        """adds a line that includes a highlight to the document"""
        block = 0
        p = self.doc.add_paragraph()
        while self._markup.WORD_HIGHLIGHT in line:
            start = line.index(self._markup.WORD_HIGHLIGHT)
            end = line.index(self._markup.WORD_HIGHLIGHT, start + 1)
            front = line[0:start]
            mid = line[start + 1 : end]
            back = line[end + 1 :]
            line = back
            run = self._add_run(p, "   " if block == 0 else "")
            run.font.name = self._fonts.BODY
            run = self._add_run(p, front)
            run.font.name = self._fonts.BODY
            run = self._add_run(p, mid)
            run.italic = True
            run.font.name = self._fonts.BODY
            if self._markup.WORD_HIGHLIGHT not in line:
                run = self._add_run(p, back)
                run.font.name = self._fonts.BODY
            block = block + 1

    def _add_run(self, p: Paragraph, text: str) -> Run:
        # print(f"add_run: {text} -- {self._metadata}")
        if self.metadata is not None:
            ws = self._get_words(text)
            # print(f"ws: {ws}")
            self.metadata.WORDS = ws + self.metadata.WORDS
            n = len(ws)
            self.metadata.WORD_COUNT = self.metadata.WORD_COUNT + n
        run = p.add_run(text)
        return run

    def _get_words(self, text: str) -> List[str]:
        words = Tokenizer.get_words(text)
        return words

    def _count_words(self, text: str) -> int:
        return len(self._get_words(text))

    def _add_jump(self, p: Paragraph) -> Run:
        #
        # TODO: get this text from config
        #
        run = p.add_run("*                   *                   *")
        return run

    def _append_line(self, lines: List[str], line: str, line_number: int):
        try:
            line = line.strip()
            #
            # blank lines are paragraph breaks. we take this time
            # to write out blocks, quotes, etc.
            #
            if line == "":
                self._append_output(lines, line, line_number)
                return
            #
            # titles
            #
            if line_number == 0:
                self.metadata.set_first_line(self._markup, line)
                return self._append_title(line)
            #
            # blocks
            #
            block = self._handle_block(line, line_number, len(lines))
            if block is None:
                # we found an escaped pipe: || meaning an
                # italicized word started a line
                line = line[1:]
            elif block:
                return
            #
            # quotes
            #
            quote = self._handle_quote(line, line_number, len(lines))
            if quote:
                return
            #
            # regular line
            #
            if self._markup.WORD_HIGHLIGHT in line:
                # if italics
                self._handle_highlights(line)
            else:
                p = self.doc.add_paragraph()
                run = self._add_run(p, f"   {line}")
                run.font.name = self._fonts.BODY
            self._last_was_break = False
        except Exception as e:
            print(f"Section failed at _append_line: error: {e}")
            traceback.print_exc()
